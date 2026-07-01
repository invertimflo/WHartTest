

from django.contrib.auth.models import User
import io

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from rest_framework import status
from rest_framework.test import APIClient

from projects.models import Project, ProjectMember
from file_management.models import FileAsset, FileManagementSetting, FileReference
from file_management.services import build_llm_attachment_context, extract_file_text
from api_interfaces.models import ApiInterface
from langgraph_integration.models import ChatSession


class FileAssetViewSetSettingsTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_superuser(username='file-admin', password='secret')
        self.project = Project.objects.create(name='File Project', creator=self.user)
        ProjectMember.objects.create(project=self.project, user=self.user, role='admin')
        self.client = APIClient(HTTP_HOST='localhost')
        self.client.force_authenticate(self.user)

    def test_file_list_does_not_conflict_with_drf_settings_property(self):
        FileAsset.objects.create(
            project=self.project,
            owner=self.user,
            original_name='demo.txt',
            extension='.txt',
            mime_type='text/plain',
            size=4,
            sha256='abc',
        )
        response = self.client.get(f'/api/projects/{self.project.id}/files/?page=1&page_size=20&ordering=-created_at')
        self.assertEqual(response.status_code, status.HTTP_200_OK)

    def test_file_settings_get_and_post_use_settings_url_without_shadowing_drf_settings(self):
        get_response = self.client.get(f'/api/projects/{self.project.id}/files/settings/')
        self.assertEqual(get_response.status_code, status.HTTP_200_OK)
        self.assertFalse(get_response.data['auto_delete_on_unbind'])
        self.assertFalse(get_response.data['auto_delete_zero_refs'])

        post_response = self.client.post(
            f'/api/projects/{self.project.id}/files/settings/',
            {'auto_delete_on_unbind': True, 'auto_delete_zero_refs': True},
            format='json',
        )
        self.assertEqual(post_response.status_code, status.HTTP_200_OK)
        self.assertTrue(post_response.data['auto_delete_on_unbind'])
        self.assertTrue(post_response.data['auto_delete_zero_refs'])
        setting = FileManagementSetting.objects.get(project=self.project)
        self.assertTrue(setting.auto_delete_on_unbind)
        self.assertTrue(setting.auto_delete_zero_refs)

    def test_file_references_endpoint_returns_object_detail(self):
        asset = FileAsset.objects.create(
            project=self.project,
            owner=self.user,
            original_name='used-by-interface.txt',
            extension='.txt',
            mime_type='text/plain',
            size=4,
            sha256='used',
        )
        interface = ApiInterface.objects.create(
            project=self.project,
            created_by=self.user,
            name='Upload Avatar',
            type=ApiInterface.TYPE_HTTP,
            method='POST',
            url='/upload',
        )
        FileReference.objects.create(
            file=asset,
            project=self.project,
            ref_type=FileReference.REF_API_INTERFACE,
            ref_id=str(interface.id),
            created_by=self.user,
        )

        response = self.client.get(f'/api/projects/{self.project.id}/files/{asset.id}/references/')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data['count'], 1)
        row = response.data['results'][0]
        self.assertEqual(row['ref_type'], FileReference.REF_API_INTERFACE)
        self.assertEqual(row['object_id'], interface.id)
        self.assertEqual(row['object_name'], 'Upload Avatar')
        self.assertIn('/upload', row['description'])

    def test_file_references_endpoint_returns_llm_chat_title_not_session_id(self):
        asset = FileAsset.objects.create(
            project=self.project,
            owner=self.user,
            original_name='chat-context.txt',
            extension='.txt',
            mime_type='text/plain',
            size=4,
            sha256='chat-used',
        )
        chat_session = ChatSession.objects.create(
            user=self.user,
            session_id='session-user-does-not-know',
            title='登录接口需求讨论',
            project=self.project,
        )
        FileReference.objects.create(
            file=asset,
            project=self.project,
            ref_type=FileReference.REF_LLM_CHAT,
            ref_id=chat_session.session_id,
            created_by=self.user,
        )

        response = self.client.get(f'/api/projects/{self.project.id}/files/{asset.id}/references/')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        row = response.data['results'][0]
        self.assertEqual(row['ref_type'], FileReference.REF_LLM_CHAT)
        self.assertEqual(row['object_id'], chat_session.session_id)
        self.assertEqual(row['object_name'], '登录接口需求讨论')
        self.assertNotIn(chat_session.session_id, row['object_name'])
        self.assertNotIn(chat_session.session_id, row['description'])

    def test_docx_text_extraction_includes_table_cells_for_llm_context(self):
        from docx import Document

        document = Document()
        document.add_paragraph('接口测试需求')
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '字段'
        table.cell(0, 1).text = '说明'
        table.cell(1, 0).text = 'username'
        table.cell(1, 1).text = '登录用户名'
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        uploaded = SimpleUploadedFile(
            'requirement.docx',
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response = self.client.post(f'/api/projects/{self.project.id}/files/', {'files': [uploaded]}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        asset = FileAsset.objects.get(id=response.data['id'])
        content = extract_file_text(asset)

        self.assertIn('接口测试需求', content)
        self.assertIn('字段', content)
        self.assertIn('说明', content)
        self.assertIn('username', content)
        self.assertIn('登录用户名', content)
        self.assertIn('|', content)


    def test_llm_attachment_context_raises_when_text_exceeds_configured_limit(self):
        uploaded = SimpleUploadedFile('large.txt', ('A' * 12).encode('utf-8'), content_type='text/plain')
        response = self.client.post(f'/api/projects/{self.project.id}/files/', {'files': [uploaded]}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        asset = FileAsset.objects.get(id=response.data['id'])

        with self.assertRaisesMessage(Exception, '超过 LLM 附件单文件上限 5 字符'):
            build_llm_attachment_context([asset], max_chars_per_file=5)

        self.assertEqual(extract_file_text(asset), 'A' * 12)



    def test_cleanup_unreferenced_endpoint_deletes_zero_reference_files(self):
        uploaded = SimpleUploadedFile('unused.txt', b'unused', content_type='text/plain')
        response = self.client.post(f'/api/projects/{self.project.id}/files/', {'files': [uploaded]}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        file_id = response.data['id']

        cleanup_response = self.client.post(f'/api/projects/{self.project.id}/files/cleanup-unreferenced/')
        self.assertEqual(cleanup_response.status_code, status.HTTP_200_OK)
        self.assertEqual(cleanup_response.data['deleted_count'], 1)
        self.assertFalse(FileAsset.objects.filter(id=file_id).exists())

    def test_physical_file_cleanup_on_delete(self):
        uploaded = SimpleUploadedFile('temp.txt', b'content', content_type='text/plain')
        response = self.client.post(f'/api/projects/{self.project.id}/files/', {'files': [uploaded]}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        file_id = response.data['id']
        asset = FileAsset.objects.get(id=file_id)
        self.assertTrue(asset.file.storage.exists(asset.file.name))

        # Perform deletion (UnifiedResponseRenderer wraps 204 into 200)
        delete_response = self.client.delete(f'/api/projects/{self.project.id}/files/{file_id}/')
        self.assertEqual(delete_response.status_code, status.HTTP_200_OK)
        self.assertFalse(FileAsset.objects.filter(id=file_id).exists())
        self.assertFalse(asset.file.storage.exists(asset.file.name))

    def test_svg_preview_blocks_xss(self):
        uploaded = SimpleUploadedFile('evil.svg', b'<svg onload="alert(1)"></svg>', content_type='image/svg+xml')
        response = self.client.post(f'/api/projects/{self.project.id}/files/', {'files': [uploaded]}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        file_id = response.data['id']

        preview_response = self.client.get(f'/api/projects/{self.project.id}/files/{file_id}/preview/')
        self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
        # Should be text/plain rather than image/svg+xml
        self.assertEqual(preview_response.headers.get('Content-Type'), 'text/plain; charset=utf-8')
        self.assertEqual(preview_response.content, b'<svg onload="alert(1)"></svg>')

    def test_missing_physical_file_returns_404(self):
        uploaded = SimpleUploadedFile('missing.txt', b'some text', content_type='text/plain')
        response = self.client.post(f'/api/projects/{self.project.id}/files/', {'files': [uploaded]}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        file_id = response.data['id']
        asset = FileAsset.objects.get(id=file_id)
        # Physically delete the file to simulate missing file
        asset.file.storage.delete(asset.file.name)

        preview_response = self.client.get(f'/api/projects/{self.project.id}/files/{file_id}/preview/')
        self.assertEqual(preview_response.status_code, status.HTTP_404_NOT_FOUND)

        download_response = self.client.get(f'/api/projects/{self.project.id}/files/{file_id}/download/')
        self.assertEqual(download_response.status_code, status.HTTP_404_NOT_FOUND)

    def test_soft_deleted_cleanup_on_zero_references(self):
        # 1. Create settings enabling auto_delete_on_unbind and auto_delete_zero_refs
        setting, _ = FileManagementSetting.objects.get_or_create(project=self.project)
        setting.auto_delete_on_unbind = True
        setting.auto_delete_zero_refs = True
        setting.save()

        # 2. Upload file
        uploaded = SimpleUploadedFile('soft.txt', b'soft-content', content_type='text/plain')
        response = self.client.post(f'/api/projects/{self.project.id}/files/', {'files': [uploaded]}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        file_id = response.data['id']
        asset = FileAsset.objects.get(id=file_id)

        # 3. Create a reference
        interface = ApiInterface.objects.create(
            project=self.project,
            created_by=self.user,
            name='Test Interface',
            type=ApiInterface.TYPE_HTTP,
            method='GET',
            url='/test',
        )
        FileReference.objects.create(
            file=asset,
            project=self.project,
            ref_type=FileReference.REF_API_INTERFACE,
            ref_id=str(interface.id),
            created_by=self.user,
        )

        # 4. Try deleting the file. Since it is referenced, it should do a soft-delete
        delete_response = self.client.delete(f'/api/projects/{self.project.id}/files/{file_id}/')
        self.assertEqual(delete_response.status_code, status.HTTP_200_OK)
        asset.refresh_from_db()
        self.assertTrue(asset.is_deleted)
        self.assertEqual(asset.status, FileAsset.STATUS_DELETED)

        # 5. Delete the referencing interface. This should trigger cascading reference deletion,
        # which should check and physically delete the unreferenced soft-deleted asset.
        interface.delete()
        self.assertFalse(FileAsset.objects.filter(id=file_id).exists())
