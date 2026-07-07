import csv
import hashlib
import io
import json
import os
from typing import Iterable, List
from django.conf import settings
from rest_framework.exceptions import ValidationError, PermissionDenied
from .models import FileAsset, FileReference, FileManagementSetting

DEFAULT_MAX_FILE_SIZE = int(os.getenv('FILE_MANAGEMENT_MAX_FILE_SIZE', str(100 * 1024 * 1024)))
DEFAULT_MAX_BIND_COUNT = int(os.getenv('FILE_MANAGEMENT_MAX_BIND_COUNT', '5'))
DEFAULT_LLM_MAX_CHARS_PER_FILE = int(os.getenv('FILE_MANAGEMENT_LLM_MAX_CHARS_PER_FILE', '0'))
TEXT_EXTENSIONS = {'.txt', '.md', '.json', '.csv', '.log', '.yaml', '.yml', '.xml', '.html', '.css', '.js', '.ts', '.py', '.svg'}


def calculate_sha256(uploaded_file) -> str:
    sha = hashlib.sha256()
    current = uploaded_file.tell() if hasattr(uploaded_file, 'tell') else None
    for chunk in uploaded_file.chunks():
        sha.update(chunk)
    if current is not None:
        uploaded_file.seek(current)
    else:
        uploaded_file.seek(0)
    return sha.hexdigest()


def normalize_file_ids(file_ids) -> List[int]:
    if file_ids in (None, ''):
        return []
    if isinstance(file_ids, str):
        try:
            file_ids = json.loads(file_ids)
        except json.JSONDecodeError as exc:
            raise ValidationError({'file_ids': 'file_ids 必须是数组。'}) from exc
    if not isinstance(file_ids, list):
        raise ValidationError({'file_ids': 'file_ids 必须是数组。'})
    result = []
    for item in file_ids:
        try:
            value = int(item)
        except (TypeError, ValueError) as exc:
            raise ValidationError({'file_ids': f'非法 file_id: {item}'}) from exc
        if value not in result:
            result.append(value)
    max_count = getattr(settings, 'FILE_MANAGEMENT_MAX_BIND_COUNT', DEFAULT_MAX_BIND_COUNT)
    if len(result) > max_count:
        raise ValidationError({'file_ids': f'单次最多绑定 {max_count} 个附件。'})
    return result


def validate_file_ids(file_ids, project, user=None, *, require_available=True):
    ids = normalize_file_ids(file_ids)
    if not ids:
        return []
    files = list(FileAsset.objects.filter(id__in=ids, project=project, is_deleted=False))
    found = {f.id for f in files}
    missing = [fid for fid in ids if fid not in found]
    if missing:
        raise ValidationError({'file_ids': f'文件不存在、已删除或不属于当前项目: {missing}'})
    unavailable = [f.id for f in files if require_available and f.status != FileAsset.STATUS_AVAILABLE]
    if unavailable:
        raise ValidationError({'file_ids': f'文件状态不可用: {unavailable}'})
    by_id = {f.id: f for f in files}
    return [by_id[fid] for fid in ids]


def sync_file_references(file_ids, project, ref_type, ref_id, user=None):
    files = validate_file_ids(file_ids, project, user)
    ref_id = str(ref_id)
    next_ids = [f.id for f in files]
    stale_qs = FileReference.objects.filter(project=project, ref_type=ref_type, ref_id=ref_id).exclude(
        file_id__in=next_ids
    )
    removed_file_ids = list(stale_qs.values_list('file_id', flat=True))
    stale_qs.delete()
    for asset in files:
        FileReference.objects.get_or_create(
            file=asset,
            project=project,
            ref_type=ref_type,
            ref_id=ref_id,
            defaults={'created_by': user},
        )
    if removed_file_ids:
        maybe_cleanup_unreferenced_files(project, candidate_file_ids=removed_file_ids, reason='unbind')
    return files


def serialize_file_for_runtime(asset: FileAsset) -> dict:
    return {
        'id': asset.id,
        'file_id': asset.id,
        'name': asset.original_name,
        'filename': asset.original_name,
        'size': asset.size,
        'mime_type': asset.mime_type,
        'extension': asset.extension,
        'path': asset.file.path if asset.file else '',
        'url': asset.file.url if asset.file else '',
    }


def _docx_paragraph_to_markdown(paragraph) -> str:
    text = (paragraph.text or '').strip()
    if not text:
        return ''

    style_name = (paragraph.style.name if paragraph.style else '').lower()
    heading_prefixes = {
        'heading 1': '# ',
        'heading 2': '## ',
        'heading 3': '### ',
        'heading 4': '#### ',
        'heading 5': '##### ',
        'heading 6': '###### ',
    }
    for style, prefix in heading_prefixes.items():
        if style in style_name:
            return f'{prefix}{text}'
    return text


def _sanitize_docx_table_cell(text: str) -> str:
    cleaned = (text or '').replace('\r', ' ').replace('\n', ' ').replace('\t', ' ')
    cleaned = ' '.join(cleaned.split())
    return cleaned.replace('|', '\\|')


def _extract_docx_table_text(table, depth: int = 0) -> str:
    """Extract a Word table as Markdown, including nested table text."""
    if depth > 3:
        return ''

    rows = []
    for row in table.rows:
        cells = []
        processed_cells = set()
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in processed_cells:
                cells.append('')
                continue
            processed_cells.add(cell_id)

            parts = []
            for paragraph in cell.paragraphs:
                text = _docx_paragraph_to_markdown(paragraph)
                if text:
                    parts.append(text)

            for nested_table in cell.tables:
                nested = _extract_docx_table_text(nested_table, depth + 1)
                if nested:
                    parts.append(f'[嵌套表格]\n{nested}')

            cells.append(_sanitize_docx_table_cell(' '.join(parts)))

        if any(cells):
            rows.append(cells)

    if not rows:
        return ''

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [''] * (max_cols - len(row)) for row in rows]
    lines = []
    for index, row in enumerate(normalized_rows):
        lines.append('| ' + ' | '.join(row) + ' |')
        if index == 0:
            lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
    return '\n'.join(lines)


def _extract_docx_text(path: str, max_chars: int | None = None) -> str:
    from docx import Document

    doc = Document(path)
    paragraph_map = {paragraph._element: paragraph for paragraph in doc.paragraphs}
    table_map = {table._element: table for table in doc.tables}
    content_parts = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            paragraph = paragraph_map.get(element)
            if paragraph:
                text = _docx_paragraph_to_markdown(paragraph)
                if text:
                    content_parts.append(text)
        elif element.tag.endswith('tbl'):
            table = table_map.get(element)
            if table:
                table_text = _extract_docx_table_text(table)
                if table_text:
                    content_parts.append(table_text)

    text = '\n\n'.join(content_parts)
    return _limit_extracted_text(text, max_chars) if max_chars is not None else text


def _limit_extracted_text(text: str, max_chars: int) -> str:
    if max_chars <= 0 or len(text) <= max_chars:
        return text
    return (
        text[:max_chars]
        + f'\n\n[预览内容超过 {max_chars} 字符，已截断。完整内容未删除。]'
    )


def _format_excel_text(path: str) -> str:
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None)
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f'## Sheet: {sheet_name}\n' + df.to_csv(index=False))
    return '\n\n'.join(parts)


def _validate_llm_attachment_text_size(asset: FileAsset, content: str, max_chars_per_file: int) -> None:
    if max_chars_per_file <= 0 or len(content) <= max_chars_per_file:
        return
    raise ValidationError({
        'file_ids': (
            f'附件 {asset.original_name} 解析后文本长度为 {len(content)} 字符，'
            f'超过 LLM 附件单文件上限 {max_chars_per_file} 字符。'
            '系统没有截断内容；请拆分附件、精简内容，或调大 FILE_MANAGEMENT_LLM_MAX_CHARS_PER_FILE 后重试。'
        )
    })


def extract_file_text(asset: FileAsset, max_chars: int | None = None) -> str:
    ext = (asset.extension or '').lower()
    name = asset.original_name
    path = asset.file.path if asset.file else ''
    if not path or not os.path.exists(path):
        return f'[文件 {name} 不存在或无法读取]'
    try:
        if ext in TEXT_EXTENSIONS:
            with open(path, 'r', encoding='utf-8', errors='replace') as fh:
                text = fh.read() if max_chars is None else fh.read(max_chars + 1)
                return _limit_extracted_text(text, max_chars) if max_chars is not None else text
        if ext in {'.xls', '.xlsx'}:
            try:
                text = _format_excel_text(path)
                return _limit_extracted_text(text, max_chars) if max_chars is not None else text
            except Exception as exc:
                return f'[Excel 文件 {name} 解析失败: {exc}]'
        if ext == '.pdf':
            try:
                from pypdf import PdfReader
                reader = PdfReader(path)
                page_texts = []
                for page in reader.pages:
                    page_texts.append(page.extract_text() or '')
                    if max_chars is not None and sum(len(item) for item in page_texts) > max_chars:
                        break
                text = '\n'.join(page_texts)
                return _limit_extracted_text(text, max_chars) if max_chars is not None else text
            except Exception as exc:
                return f'[PDF 文件 {name} 暂无法解析: {exc}]'
        if ext == '.docx':
            try:
                return _extract_docx_text(path, max_chars)
            except Exception as exc:
                return f'[DOCX 文件 {name} 暂无法解析: {exc}]'
        return f'[附件 {name}，类型 {asset.mime_type or ext or "unknown"}，大小 {asset.size} 字节，当前仅作为文件引用使用]'
    except Exception as exc:
        return f'[文件 {name} 读取失败: {exc}]'


def build_llm_attachment_context(files: Iterable[FileAsset], max_chars_per_file: int = DEFAULT_LLM_MAX_CHARS_PER_FILE) -> str:
    parts = []
    for asset in files:
        content = extract_file_text(asset)
        _validate_llm_attachment_text_size(asset, content, max_chars_per_file)
        parts.append(
            f'\n\n--- 附件 file_id={asset.id} name={asset.original_name} size={asset.size} type={asset.mime_type} ---\n{content}'
        )
    return ''.join(parts)



def delete_file_asset_if_unreferenced(file_id, project=None) -> bool:
    """Delete a managed file only when it is not referenced anywhere.

    This is used for files uploaded as inline operation parameters, such as UI
    automation upload-file steps. It removes both the DB row and the underlying
    storage file when the file has no FileReference records.
    """
    if not file_id:
        return False
    qs = FileAsset.objects.filter(id=file_id)
    if project is not None:
        qs = qs.filter(project=project)
    asset = qs.first()
    if not asset:
        return False
    if asset.references.exists():
        return False
    storage = asset.file.storage if asset.file else None
    storage_name = asset.file.name if asset.file else ''
    asset.delete()
    if storage and storage_name:
        try:
            if storage.exists(storage_name):
                storage.delete(storage_name)
        except Exception:
            # DB deletion has already succeeded; avoid breaking the caller on a
            # best-effort physical-file cleanup failure.
            pass
    return True



def get_file_management_setting(project):
    setting, _ = FileManagementSetting.objects.get_or_create(project=project)
    return setting


def cleanup_unreferenced_files(project, *, only_file_ids=None) -> int:
    """Delete unreferenced managed files for a project.

    When only_file_ids is supplied, cleanup is limited to those files. Physical
    storage deletion is best-effort via delete_file_asset_if_unreferenced().
    """
    qs = FileAsset.objects.filter(project=project)
    if only_file_ids is not None:
        qs = qs.filter(id__in=list(only_file_ids))
    ids = list(qs.filter(references__isnull=True).values_list('id', flat=True).distinct())
    count = 0
    for fid in ids:
        if delete_file_asset_if_unreferenced(fid, project=project):
            count += 1
    return count


def maybe_cleanup_unreferenced_files(project, *, candidate_file_ids=None, reason='unbind') -> int:
    setting = get_file_management_setting(project)
    if reason == 'unbind' and not setting.auto_delete_on_unbind:
        return 0
    if reason == 'zero_refs' and not setting.auto_delete_zero_refs:
        return 0
    if setting.auto_delete_zero_refs and candidate_file_ids is None:
        return cleanup_unreferenced_files(project)
    return cleanup_unreferenced_files(project, only_file_ids=candidate_file_ids or [])



def resolve_file_reference_detail(ref: FileReference) -> dict:
    """Return a user-friendly description for a FileReference."""
    ref_id = str(ref.ref_id)
    data = {
        'id': ref.id,
        'file_id': ref.file_id,
        'project': ref.project_id,
        'ref_type': ref.ref_type,
        'ref_type_label': dict(FileReference.REF_CHOICES).get(ref.ref_type, ref.ref_type),
        'ref_id': ref_id,
        'object_id': ref_id,
        'object_name': '',
        'description': '',
        'created_by': getattr(ref.created_by, 'username', '') if ref.created_by_id else '',
        'created_at': ref.created_at,
    }
    try:
        if ref.ref_type == FileReference.REF_API_INTERFACE:
            from api_interfaces.models import ApiInterface
            obj = ApiInterface.objects.select_related('module').filter(id=ref_id, project_id=ref.project_id).first()
            if obj:
                data.update({
                    'object_id': obj.id,
                    'object_name': obj.name,
                    'description': f"{obj.get_type_display() if hasattr(obj, 'get_type_display') else obj.type} {obj.method or obj.sql_method or ''} {obj.url or obj.sql or ''}".strip(),
                    'module_name': obj.module.name if obj.module_id else '',
                })
        elif ref.ref_type == FileReference.REF_API_TESTCASE:
            from api_testcases.models import ApiTestCase
            obj = ApiTestCase.objects.select_related('group').filter(id=ref_id, project_id=ref.project_id).first()
            if obj:
                data.update({
                    'object_id': obj.id,
                    'object_name': obj.name,
                    'description': obj.description or '',
                    'module_name': obj.group.name if obj.group_id else '',
                })
        elif ref.ref_type == FileReference.REF_API_INTERFACE_CASE:
            from api_testcases.models import ApiInterfaceCase
            obj = ApiInterfaceCase.objects.select_related('group', 'interface').filter(id=ref_id, project_id=ref.project_id).first()
            if obj:
                data.update({
                    'object_id': obj.id,
                    'object_name': obj.name,
                    'description': obj.description or '',
                    'module_name': obj.interface.module.name if obj.interface and obj.interface.module_id else '',
                    'interface_name': obj.interface.name if obj.interface_id else '',
                })
        elif ref.ref_type == FileReference.REF_UI_TESTCASE:
            from ui_automation.models import UiTestCase
            obj = UiTestCase.objects.select_related('module').filter(id=ref_id, project_id=ref.project_id).first()
            if obj:
                data.update({
                    'object_id': obj.id,
                    'object_name': obj.name,
                    'description': obj.description or '',
                    'module_name': obj.module.name if obj.module_id else '',
                })
        elif ref.ref_type == FileReference.REF_UI_PAGE_STEPS:
            from ui_automation.models import UiPageSteps, UiPageStepsDetailed
            if ref_id.startswith('detail:'):
                detail_id = ref_id.split(':', 1)[1]
                detail = UiPageStepsDetailed.objects.select_related(
                    'page_step', 'page_step__page', 'page_step__module', 'element'
                ).filter(id=detail_id, page_step__project_id=ref.project_id).first()
                if detail:
                    page_step = detail.page_step
                    data.update({
                        'object_id': detail.id,
                        'object_name': f"{page_step.name} / Step {detail.step_sort + 1} / {detail.ope_key or ''}".strip(),
                        'description': detail.description or (detail.element.name if detail.element_id else ''),
                        'parent_id': page_step.id,
                        'parent_name': page_step.name,
                        'page_name': page_step.page.name if page_step.page_id else '',
                        'module_name': page_step.module.name if page_step.module_id else '',
                    })
            else:
                obj = UiPageSteps.objects.select_related('page', 'module').filter(id=ref_id, project_id=ref.project_id).first()
                if obj:
                    data.update({
                        'object_id': obj.id,
                        'object_name': obj.name,
                        'description': obj.description or '',
                        'page_name': obj.page.name if obj.page_id else '',
                        'module_name': obj.module.name if obj.module_id else '',
                    })
        elif ref.ref_type == FileReference.REF_LLM_CHAT:
            from langgraph_integration.models import ChatSession
            chat_session = ChatSession.objects.filter(
                session_id=ref_id,
                project_id=ref.project_id,
            ).first()
            if chat_session:
                data.update({
                    'object_id': chat_session.session_id,
                    'object_name': chat_session.title or '未命名对话',
                    'description': '',
                })
            else:
                data.update({'object_name': '已删除或不可用的对话', 'description': ''})
    except Exception as exc:
        data['description'] = f'引用详情解析失败: {exc}'
    if not data.get('object_name'):
        data['object_name'] = ref_id
    return data
